"""
Document Export Service - Converts LaTeX resume content to other formats.
All conversions are rule-based (no LLM), synchronous, and fast (<200ms).
"""
import io
import logging
import re
from typing import Any, Dict, List

logger = logging.getLogger(__name__)


class DocumentExportService:
    """Service for exporting LaTeX resumes to other file formats."""

    # ─── Text Extraction ────────────────────────────────────────────────────

    def to_text(self, latex_content: str) -> str:
        """Convert LaTeX to plain text, preserving structure via newlines."""
        # Convert through markdown (which handles LaTeX → structured text),
        # then strip markdown formatting markers.
        md = self.to_markdown(latex_content)

        # Strip markdown heading markers (## → UPPERCASE line)
        def heading_replacer(m):
            level = len(m.group(1))
            title = m.group(2).strip()
            if level == 2:
                return f'\n{title.upper()}\n{"─" * len(title)}'
            return f'\n{title}'

        text = re.sub(r'^(#{1,4})\s+(.+)$', heading_replacer, md, flags=re.MULTILINE)

        # Strip bold/italic markers but keep text
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'<u>(.+?)</u>', r'\1', text)

        # Markdown links → text only
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        text = re.sub(r'<(https?://[^>]+)>', r'\1', text)

        # List bullets → •
        text = re.sub(r'^-\s+', '• ', text, flags=re.MULTILINE)

        # Collapse 3+ blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)

        return text.strip()

    # ─── Markdown ───────────────────────────────────────────────────────────

    def to_markdown(self, latex_content: str) -> str:
        """Convert LaTeX resume to Markdown."""
        text = latex_content

        # Escaped special characters → plain text (before any other processing)
        text = text.replace(r'\%', '%')
        text = text.replace(r'\&', '&')
        text = text.replace(r'\$', '$')
        text = text.replace(r'\#', '#')
        text = text.replace(r'\_', '_')
        text = text.replace(r'\~', '~')
        text = text.replace(r'\^', '^')

        # Remove font size commands inside textbf/textit before processing formatting
        text = re.sub(r'\\(?:Large|large|LARGE|huge|Huge|small|footnotesize|normalsize)\b\s*', '', text)

        # Section headings (must be before generic command removal)
        text = re.sub(r'\\section\*?\{([^}]*)\}', r'\n## \1\n', text)
        text = re.sub(r'\\subsection\*?\{([^}]*)\}', r'\n### \1\n', text)
        text = re.sub(r'\\subsubsection\*?\{([^}]*)\}', r'\n#### \1\n', text)

        # Text formatting (must be before generic removal)
        # Strip inner whitespace so **  Jane  ** → **Jane**
        text = re.sub(r'\\textbf\{\s*([^}]*?)\s*\}', lambda m: f'**{m.group(1).strip()}**', text)
        text = re.sub(r'\\textit\{\s*([^}]*?)\s*\}', lambda m: f'*{m.group(1).strip()}*', text)
        text = re.sub(r'\\emph\{\s*([^}]*?)\s*\}', lambda m: f'*{m.group(1).strip()}*', text)
        text = re.sub(r'\\underline\{\s*([^}]*?)\s*\}', lambda m: f'<u>{m.group(1).strip()}</u>', text)
        text = re.sub(r'\\texttt\{\s*([^}]*?)\s*\}', lambda m: f'`{m.group(1).strip()}`', text)

        # Links (before generic removal)
        text = re.sub(r'\\href\{([^}]*)\}\{([^}]*)\}', r'[\2](\1)', text)
        text = re.sub(r'\\url\{([^}]*)\}', r'<\1>', text)

        # Lists
        text = re.sub(r'\\begin\{itemize\}', '', text)
        text = re.sub(r'\\end\{itemize\}', '', text)
        text = re.sub(r'\\begin\{enumerate\}', '', text)
        text = re.sub(r'\\end\{enumerate\}', '', text)
        text = re.sub(r'[ \t]*\\item\s+', '- ', text)

        # LaTeX line breaks (\\) → actual newline
        text = re.sub(r'\\\\', '\n', text)

        # Horizontal fill → separator, spacing commands → remove
        text = re.sub(r'\\hfill\b', '  —  ', text)
        text = re.sub(r'\\vspace\*?\{[^}]*\}', '', text)
        text = re.sub(r'\\hspace\*?\{[^}]*\}', ' ', text)
        text = re.sub(r'\\noindent\b', '', text)
        text = re.sub(r'\\centering\b', '', text)

        # Em-dash and en-dash
        text = text.replace('---', '—')
        text = text.replace('--', '–')

        # Remove document structure
        text = re.sub(r'\\begin\{document\}', '', text)
        text = re.sub(r'\\end\{document\}', '', text)
        text = re.sub(r'\\documentclass[^\n]*\n?', '', text)
        text = re.sub(r'\\usepackage[^\n]*\n?', '', text)
        text = re.sub(r'\\geometry[^\n]*\n?', '', text)

        # Remove begin/end environments
        text = re.sub(r'\\begin\{[^}]*\}', '', text)
        text = re.sub(r'\\end\{[^}]*\}', '', text)

        # Remove LaTeX command names but KEEP their brace-enclosed arguments.
        # e.g. \resumeSubheading{Company}{Role} → "Company Role"
        # Pattern: strip \cmdname and optional [...] arg, but leave the {...} content.
        text = re.sub(r'\\[a-zA-Z]+\*?(?:\[[^\]]*\])?', '', text)
        # Remove lone backslashes
        text = re.sub(r'\\', '', text)

        # Clean up braces, extra whitespace
        text = re.sub(r'[{}]', '', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        # Clean trailing spaces on lines
        text = re.sub(r' +\n', '\n', text)

        return text.strip()

    # ─── HTML ────────────────────────────────────────────────────────────────

    def to_html(self, latex_content: str) -> str:
        """Convert LaTeX resume to HTML."""
        try:
            import mistune
            md_content = self.to_markdown(latex_content)
            body = mistune.html(md_content)
        except ImportError:
            # Fallback: basic HTML from text — escape to prevent XSS
            import html as _html
            text = self.to_text(latex_content)
            paragraphs = [f'<p>{_html.escape(p)}</p>' for p in text.split('\n\n') if p.strip()]
            body = '\n'.join(paragraphs)

        return (
            '<!DOCTYPE html>\n'
            '<html lang="en">\n'
            '<head>\n'
            '  <meta charset="UTF-8">\n'
            '  <meta name="viewport" content="width=device-width, initial-scale=1.0">\n'
            '  <title>Resume</title>\n'
            '  <style>\n'
            '    body { font-family: Arial, sans-serif; max-width: 800px; margin: 2rem auto; '
            'padding: 0 1rem; line-height: 1.6; color: #333; }\n'
            '    h2 { border-bottom: 1px solid #ccc; padding-bottom: 0.3rem; margin-top: 2rem; }\n'
            '    ul { margin: 0.5rem 0; padding-left: 1.5rem; }\n'
            '    a { color: #0066cc; }\n'
            '  </style>\n'
            '</head>\n'
            '<body>\n'
            f'{body}\n'
            '</body>\n'
            '</html>'
        )

    # ─── JSON Resume ─────────────────────────────────────────────────────────

    def _to_plain_text(self, latex_content: str) -> str:
        """Internal: plain text without decorative section rulers, for structured parsing."""
        md = self.to_markdown(latex_content)
        # Strip markdown markers, keep text
        text = re.sub(r'^#{1,4}\s+(.+)$', r'\1', md, flags=re.MULTILINE)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'<u>(.+?)</u>', r'\1', text)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        text = re.sub(r'^-\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def to_json(self, latex_content: str) -> Dict[str, Any]:
        """Convert LaTeX resume to JSON Resume schema (https://jsonresume.org/schema)."""
        text = self._to_plain_text(latex_content)
        lines = [line.strip() for line in text.split('\n') if line.strip()]

        # Extract basic contact info
        import re as _re
        email = ''
        phone = ''
        name = lines[0] if lines else ''

        email_match = _re.search(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}', text)
        if email_match:
            email = email_match.group(0)

        phone_match = _re.search(r'(?:\+?1[-.\s]?)?(?:\(?[2-9]\d{2}\)?[-.\s]?)[2-9]\d{2}[-.\s]?\d{4}', text)
        if phone_match:
            phone = phone_match.group(0)

        linkedin_match = _re.search(r'linkedin\.com/in/([a-zA-Z0-9\-]+)', text)
        linkedin_url = f"https://linkedin.com/in/{linkedin_match.group(1)}" if linkedin_match else ''

        github_match = _re.search(r'github\.com/([a-zA-Z0-9\-]+)', text)
        github_url = f"https://github.com/{github_match.group(1)}" if github_match else ''

        # Build section-based content
        SECTION_RE = _re.compile(
            r'^(?:EXPERIENCE|WORK|EDUCATION|SKILLS|PROJECTS|CERTIFICATIONS|'
            r'SUMMARY|OBJECTIVE|AWARDS|PUBLICATIONS)\b',
            _re.IGNORECASE
        )
        current_section = None
        sections: Dict[str, List[str]] = {}
        for line in lines:
            if SECTION_RE.match(line):
                current_section = line.upper().split()[0]
                sections[current_section] = []
            elif current_section:
                sections[current_section].append(line)

        # Skills extraction
        skills_list = []
        if 'SKILLS' in sections:
            skills_text = ' '.join(sections['SKILLS'])
            skills_list = [s.strip() for s in _re.split(r'[,;|•·]', skills_text) if s.strip()]

        result = {
            "$schema": "https://raw.githubusercontent.com/jsonresume/resume-schema/v1.0.0/schema.json",
            "basics": {
                "name": name,
                "email": email,
                "phone": phone,
                "url": "",
                "summary": ' '.join(sections.get('SUMMARY', sections.get('OBJECTIVE', []))),
                "profiles": [],
            },
            "work": [],
            "education": [],
            "skills": [{"name": s} for s in skills_list[:20]],
            "projects": [],
            "awards": [],
            "certificates": [],
        }

        if linkedin_url:
            result["basics"]["profiles"].append({"network": "LinkedIn", "url": linkedin_url})
        if github_url:
            result["basics"]["profiles"].append({"network": "GitHub", "url": github_url})

        return result

    # ─── YAML ────────────────────────────────────────────────────────────────

    def to_yaml(self, latex_content: str) -> str:
        """Convert LaTeX resume to YAML (JSON Resume schema)."""
        import yaml
        data = self.to_json(latex_content)
        return yaml.dump(data, allow_unicode=True, sort_keys=False, default_flow_style=False)

    # ─── XML ─────────────────────────────────────────────────────────────────

    def to_xml(self, latex_content: str) -> str:
        """Convert LaTeX resume to XML."""
        data = self.to_json(latex_content)
        lines = ['<?xml version="1.0" encoding="UTF-8"?>', '<resume>']

        def dict_to_xml(d: Any, indent: int = 2) -> List[str]:
            result = []
            pad = ' ' * indent
            if isinstance(d, dict):
                for k, v in d.items():
                    tag = re.sub(r'[^a-zA-Z0-9_\-]', '_', str(k))
                    if isinstance(v, (dict, list)):
                        result.append(f'{pad}<{tag}>')
                        result.extend(dict_to_xml(v, indent + 2))
                        result.append(f'{pad}</{tag}>')
                    else:
                        safe_v = str(v).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        result.append(f'{pad}<{tag}>{safe_v}</{tag}>')
            elif isinstance(d, list):
                for i, item in enumerate(d):
                    result.append(f'{pad}<item>')
                    result.extend(dict_to_xml(item, indent + 2))
                    result.append(f'{pad}</item>')
            return result

        lines.extend(dict_to_xml(data))
        lines.append('</resume>')
        return '\n'.join(lines)

    # ─── DOCX ────────────────────────────────────────────────────────────────

    def to_docx(self, latex_content: str) -> bytes:
        """Convert LaTeX resume to DOCX format."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
        except ImportError:
            raise ValueError("python-docx not installed. Run: pip install python-docx")

        doc = Document()

        # Set page margins
        for section in doc.sections:
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)

        # Remove default empty paragraph
        for para in doc.paragraphs:
            p = para._element
            p.getparent().remove(p)

        # Extract content structure from LaTeX
        md_content = self.to_markdown(latex_content)
        lines = md_content.split('\n')

        for line in lines:
            line = line.rstrip()
            if not line:
                continue
            elif line.startswith('## '):
                # Section heading
                heading_text = line[3:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                run = p.add_run(heading_text.upper())
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
                # Add underline rule
                p.paragraph_format.border_bottom = True if hasattr(p.paragraph_format, 'border_bottom') else None
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            elif line.startswith('### '):
                p = doc.add_paragraph()
                run = p.add_run(line[4:].strip())
                run.bold = True
                run.font.size = Pt(11)
            elif line.startswith('- '):
                # Bullet point
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:].strip())
                p.paragraph_format.space_after = Pt(2)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.bold = True
            else:
                p = doc.add_paragraph()
                p.add_run(line)
                p.paragraph_format.space_after = Pt(2)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


document_export_service = DocumentExportService()
